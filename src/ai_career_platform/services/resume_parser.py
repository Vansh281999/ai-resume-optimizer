import io
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)

from ai_career_platform.ai_providers.factory import get_multi_provider


class ResumeIngestionError(RuntimeError):
    pass


class ResumeFileValidationError(ResumeIngestionError):
    pass


CANONICAL_SKILLS = {
    # Programming Languages
    "python": "programming_languages", "java": "programming_languages", "c++": "programming_languages", "c#": "programming_languages",
    "javascript": "programming_languages", "typescript": "programming_languages", "go": "programming_languages", "golang": "programming_languages",
    "rust": "programming_languages", "ruby": "programming_languages", "php": "programming_languages", "swift": "programming_languages",
    "kotlin": "programming_languages", "scala": "programming_languages", "r": "programming_languages", "matlab": "programming_languages",
    "perl": "programming_languages", "lua": "programming_languages", "dart": "programming_languages", "objective-c": "programming_languages",
    "sql": "programming_languages", "html": "programming_languages", "css": "programming_languages", "shell": "programming_languages",
    "bash": "programming_languages", "powershell": "programming_languages", "c": "programming_languages", "assembly": "programming_languages",
    "haskell": "programming_languages", "erlang": "programming_languages", "clojure": "programming_languages", "elixir": "programming_languages",
    "f#": "programming_languages", "groovy": "programming_languages", "julia": "programming_languages", "sas": "programming_languages",
    "awk": "programming_languages", "sed": "programming_languages",
    # Frameworks & Libraries (explicit entries for proper categorization)
    "react": "frameworks", "react.js": "frameworks", "next.js": "frameworks", "nextjs": "frameworks",
    "angular": "frameworks", "vue": "frameworks", "vue.js": "frameworks", "nuxt": "frameworks", "nuxt.js": "frameworks",
    "svelte": "frameworks", "sveltekit": "frameworks", "django": "frameworks", "flask": "frameworks", "fastapi": "frameworks",
    "spring": "frameworks", "spring boot": "frameworks", "express": "frameworks", "express.js": "frameworks", "nestjs": "frameworks",
    "asp.net": "frameworks", ".net": "frameworks", "laravel": "frameworks", "rails": "frameworks", "ruby on rails": "frameworks",
    "gatsby": "frameworks", "ember": "frameworks", "backbone": "frameworks", "backbone.js": "frameworks", "jquery": "frameworks",
    "bootstrap": "frameworks", "tailwind": "frameworks", "tailwind css": "frameworks", "tailwindcss": "frameworks",
    "material-ui": "frameworks", "chakra": "frameworks", "chakra ui": "frameworks", "redux": "frameworks", "mobx": "frameworks",
    "graphql": "frameworks", "rest": "frameworks", "rest api": "frameworks", "node.js": "frameworks", "nodejs": "frameworks",
    "jest": "frameworks", "mocha": "frameworks", "pytest": "frameworks", "junit": "frameworks", "cypress": "frameworks",
    "selenium": "frameworks", "puppeteer": "frameworks", "playwright": "frameworks", "storybook": "frameworks",
    "three.js": "frameworks", "d3.js": "frameworks", "chart.js": "frameworks", "socket.io": "frameworks",
    # AI/ML Technologies
    "tensorflow": "ai_ml_technologies", "pytorch": "ai_ml_technologies", "keras": "ai_ml_technologies",
    "scikit-learn": "ai_ml_technologies", "sklearn": "ai_ml_technologies", "pandas": "ai_ml_technologies",
    "numpy": "ai_ml_technologies", "opencv": "ai_ml_technologies", "huggingface": "ai_ml_technologies",
    "langchain": "ai_ml_technologies", "llm": "ai_ml_technologies", "gpt": "ai_ml_technologies", "bert": "ai_ml_technologies",
    "transformers": "ai_ml_technologies", "spacy": "ai_ml_technologies", "nltk": "ai_ml_technologies",
    "matplotlib": "ai_ml_technologies", "seaborn": "ai_ml_technologies", "plotly": "ai_ml_technologies",
    "jupyter": "ai_ml_technologies", "colab": "ai_ml_technologies", "machine learning": "ai_ml_technologies",
    "deep learning": "ai_ml_technologies", "nlp": "ai_ml_technologies", "computer vision": "ai_ml_technologies",
    "data science": "ai_ml_technologies", "data analysis": "ai_ml_technologies", "prisma": "databases",
    # Databases
    "postgresql": "databases", "postgres": "databases", "mysql": "databases", "mongodb": "databases", "mongo": "databases",
    "redis": "databases", "redis cache": "databases", "elasticsearch": "databases", "cassandra": "databases",
    "dynamodb": "databases", "sqlite": "databases", "oracle": "databases", "oracle db": "databases",
    "mariadb": "databases", "neo4j": "databases", "couchdb": "databases", "firestore": "databases",
    "supabase": "databases", "sqlalchemy": "databases", "sequelize": "databases", "mongoose": "databases",
    "django orm": "databases", "mongodb": "databases",
    # Cloud
    "aws": "cloud_technologies", "azure": "cloud_technologies", "gcp": "cloud_technologies", "google cloud": "cloud_technologies",
    "heroku": "cloud_technologies", "vercel": "cloud_technologies", "netlify": "cloud_technologies",
    "digitalocean": "cloud_technologies", "linode": "cloud_technologies", "cloudflare": "cloud_technologies",
    "aws ec2": "cloud_technologies", "aws s3": "cloud_technologies", "aws lambda": "cloud_technologies",
    # DevOps
    "docker": "devops_tools", "kubernetes": "devops_tools", "k8s": "devops_tools", "terraform": "devops_tools",
    "ansible": "devops_tools", "jenkins": "devops_tools", "git": "devops_tools", "gitlab": "devops_tools",
    "bitbucket": "devops_tools", "ci/cd": "devops_tools", "cicd": "devops_tools", "linux": "devops_tools",
    "nginx": "devops_tools", "apache": "devops_tools", "prometheus": "devops_tools", "grafana": "devops_tools",
    "github actions": "devops_tools", "gitlab ci": "devops_tools", "circleci": "devops_tools", "travis ci": "devops_tools",
    "argo cd": "devops_tools", "helm": "devops_tools", "istio": "devops_tools", "envoy": "devops_tools",
    "vagrant": "devops_tools", "packer": "devops_tools",
    # Soft Skills
    "leadership": "soft_skills", "communication": "soft_skills", "teamwork": "soft_skills",
    "problem solving": "soft_skills", "critical thinking": "soft_skills", "agile": "soft_skills",
    "scrum": "soft_skills", "project management": "soft_skills", "time management": "soft_skills",
    "adaptability": "soft_skills", "creativity": "soft_skills", "collaboration": "soft_skills",
    "mentoring": "soft_skills", "public speaking": "soft_skills", "negotiation": "soft_skills",
    "stakeholder management": "soft_skills", "conflict resolution": "soft_skills",
    "strategic planning": "soft_skills", "team building": "soft_skills",
}

INDIAN_CITIES = {
    "bangalore", "bengaluru", "mumbai", "bombay", "delhi", "new delhi", "hyderabad", "chennai", "madras",
    "kolkata", "calcutta", "pune", "ahmedabad", "jaipur", "surat", "lucknow", "kanpur", "nagpur",
    "indore", "thane", "bhopal", "visakhapatnam", "patna", "vadodara", "ludhiana", "agra", "nashik",
    "faridabad", "meerut", "rajkot", "varanasi", "srinagar", "aurangabad", "dhanbad", "amritsar",
    "navi mumbai", "prayagraj", "ranchi", "howrah", "coimbatore", "jabalpur", "gwalior", "vijayawada",
    "jodhpur", "madurai", "raipur", "kota", "guwahati", "chandigarh", "solapur", "hubli", "dharwad",
    "bhubaneswar", "trivandrum", "thiruvananthapuram", "kochi", "cochin", "mysore", "mysuru",
    "vellore", "salem", "tiruchirappalli", "trichy", "coimbatore", "guntur", "bhavnagar", "dehradun",
    "jamshedpur", "asansol", "allahabad", "gorakhpur", "aligarh", "bareilly", "moradabad",
    "gurgaon", "gurugram", "noida", "faridabad", "ghaziabad", "thiruvananthapuram",
    "vit vellore", "iit bombay", "iit delhi", "iit madras", "iit kharagpur", "iit kanpur",
    "bits pilani", "bits goa", "bits hyderabad", "nit trichy", "nit warangal", "nit surat",
}

US_CITIES = {
    "new york", "los angeles", "chicago", "houston", "phoenix", "philadelphia", "san antonio",
    "san diego", "dallas", "san jose", "austin", "jacksonville", "fort worth", "columbus",
    "charlotte", "san francisco", "indianapolis", "seattle", "denver", "washington", "boston",
    "el paso", "nashville", "detroit", "portland", "memphis", "oklahoma city", "las vegas",
    "louisville", "baltimore", "milwaukee", "albuquerque", "tucson", "fresno", "sacramento",
    "mesa", "atlanta", "kansas city", "colorado springs", "raleigh", "omaha", "miami",
    "long beach", "virginia beach", "oakland", "minneapolis", "tulsa", "tampa", "arlington",
    "new orleans", "wichita", "cleveland", "bakersfield", "tacoma", "aurora", "anaheim",
    "santa ana", "st louis", "riverside", "corpus christi", "lexington", "pittsburgh",
    "anchorage", "stockton", "cincinnati", "st. paul", "toledo", "newark", "chandler",
    "las vegas", "boulder", "palo alto", "mountain view", "sunnyvale", "santa clara",
    "seattle", "redmond", "bellevue", "kirkland", "everett", "tacoma", "olympia",
    "remote", "hybrid", "work from home", "wfh", "work-from-home",
}


class SectionClassifier:
    SECTION_PATTERNS = {
        "header": [r"^[A-Za-z][A-Za-z\s]{2,}$"],
        "summary": [
            r"^\s*(summary|professional[\s_]summary|executive[\s_]summary|profile|about[\s_]me|career[\s_]summary|personal[\s_]profile)\s*:?\s*$",
        ],
        "education": [
            r"^\s*(education|academic|academics|qualification|qualifications|educational[\s_]background)\s*:?\s*$",
        ],
        "experience": [
            r"^\s*(experience|work[\s_]experience|professional[\s_]experience|employment[\s_]history|work[\s_]history|professional[\s_]history)\s*:?\s*$",
        ],
        "internships": [
            r"^\s*(internship|internships)\s*:?\s*$",
        ],
        "projects": [
            r"^\s*(projects?|personal[\s_]projects?|portfolio|academic[\s_]projects?|key[\s_]projects?)\s*:?\s*$",
        ],
        "skills": [
            r"^\s*(skills?|technical[\s_]skills?|core[\s_]competenc(?:y|ies)|technologies|programming[\s_]languages|tools?\s*&?\s*technologies)\s*:?\s*$",
        ],
        "certifications": [
            r"^\s*(certifications?|certificates?|certified|credentials?|licenses?)\s*:?\s*$",
        ],
        "achievements": [
            r"^\s*(achievements?|awards?|honors?|recognitions?|accomplishments?)\s*:?\s*$",
        ],
    }

    # Any line matching these patterns marks an immediate section boundary
    BOUNDARY_PATTERNS = [
        r"^\s*(summary|professional[\s_]summary|executive[\s_]summary|profile|about[\s_]me|career[\s_]summary|personal[\s_]profile)\s*:?\s*$",
        r"^\s*(education|academic|academics|qualification|qualifications|educational[\s_]background)\s*:?\s*$",
        r"^\s*(experience|work[\s_]experience|professional[\s_]experience|employment[\s_]history|work[\s_]history|professional[\s_]history)\s*:?\s*$",
        r"^\s*(internship|internships|internships?)\s*:?\s*$",
        r"^\s*(projects?|personal[\s_]projects?|portfolio|academic[\s_]projects?|key[\s_]projects?)\s*:?\s*$",
        r"^\s*(skills?|technical[\s_]skills?|core[\s_]competenc(?:y|ies)|technologies|programming[\s_]languages|tools?\s*&?\s*technologies)\s*:?\s*$",
        r"^\s*(certifications?|certificates?|certified|credentials?|licenses?)\s*:?\s*$",
        r"^\s*(achievements?|awards?|honors?|recognitions?|accomplishments?)\s*:?\s*$",
        r"^\s*(references?)\s*:?\s*$",
        r"^\s*(declaration)\s*:?\s*$",
    ]

    SECTION_BOUNDARIES = re.compile(
        r"^\s*(summary|professional[\s_]summary|executive[\s_]summary|profile|about[\s_]me|career[\s_]summary|personal[\s_]profile|"
        r"education|academic|academics|qualification|qualifications|educational[\s_]background|"
        r"experience|work[\s_]experience|professional[\s_]experience|employment[\s_]history|work[\s_]history|professional[\s_]history|"
        r"internship|internships|"
        r"projects?|personal[\s_]projects?|portfolio|academic[\s_]projects?|key[\s_]projects?|"
        r"skills?|technical[\s_]skills?|core[\s_]competenc(?:y|ies)|technologies|programming[\s_]languages|tools?\s*&?\s*technologies|"
        r"certifications?|certificates?|certified|credentials?|licenses?|"
        r"achievements?|awards?|honors?|recognitions?|accomplishments?|"
        r"references?|declaration)\s*:?\s*$",
        re.IGNORECASE
    )

    @staticmethod
    def is_boundary(line: str) -> bool:
        stripped = line.strip().lower()
        for pattern in SectionClassifier.BOUNDARY_PATTERNS:
            if re.match(pattern, stripped, re.IGNORECASE):
                return True
        if stripped in {"", "---", "***", "===", "___"}:
            return True
        return False

    @staticmethod
    def classify(lines: List[str]) -> Dict[str, List[int]]:
        sections: Dict[str, List[int]] = {name: [] for name in SectionClassifier.SECTION_PATTERNS}
        current_section = "header"
        header_end = min(5, len(lines))

        for idx, line in enumerate(lines):
            if SectionClassifier.SECTION_BOUNDARIES.match(line.strip()):
                matched = False
                for section_name, patterns in SectionClassifier.SECTION_PATTERNS.items():
                    for pattern in patterns:
                        if pattern != r"^[A-Za-z][A-Za-z\s]{2,}$" and re.match(pattern, line.strip(), re.IGNORECASE):
                            current_section = section_name
                            matched = True
                            break
                    if matched:
                        break
            else:
                sections[current_section].append(idx)

        if "header" in sections and sections["header"]:
            sections["header"] = [i for i in sections["header"] if i < header_end]

        return sections


class ConfidenceScorer:
    HIGH = 0.9
    MEDIUM = 0.7
    LOW = 0.5
    VERY_LOW = 0.3

    @staticmethod
    def scored(value: Any, confidence: float) -> Dict[str, Any]:
        return {"value": value if value is not None else "", "confidence": max(0.0, min(1.0, confidence))}

    @classmethod
    def from_match(cls, value: str, pattern_strength: str = "medium") -> Dict[str, Any]:
        if pattern_strength == "high":
            return cls.scored(value, cls.HIGH)
        if pattern_strength == "medium":
            return cls.scored(value, cls.MEDIUM)
        return cls.scored(value, cls.LOW)

    @classmethod
    def best(cls, candidates: List[Dict[str, Any]]) -> Dict[str, Any]:
        if not candidates:
            return cls.scored("", cls.VERY_LOW)
        return max(candidates, key=lambda c: c.get("confidence", 0))


class ResumeIngestionPipeline:
    allowed_extensions = {".pdf", ".docx", ".doc", ".txt"}
    allowed_mime_types = {
        ".pdf": {"application/pdf"},
        ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
        ".doc": {"application/msword", "application/vnd.ms-word"},
        ".txt": {"text/plain", "text/markdown", "application/octet-stream", ""},
    }

    def __init__(self, providers: Optional[List[str]] = None, model: str = "gemini-2.0-flash", max_text_chars: int = 30000):
        self.model = model
        self.max_text_chars = max_text_chars
        self.provider = get_multi_provider(providers=providers or ["gemini", "openrouter", "ollama"], model=model)
        self.last_extraction_metadata: Dict[str, Any] = {}
        self.system_prompt = (
            "You are an expert resume parser. Extract all information from the resume and return ONLY valid JSON. "
            "Do not include markdown, explanations, or comments. If a field is not present, use an empty string, null, or an empty array.\n\n"
            "Schema:\n"
            "{\n"
            '  "personal_info": {"full_name": "", "email": "", "phone": "", "location": "", "linkedin_url": "", "github_url": "", "portfolio_url": "", "leetcode_url": "", "hackerrank_url": "", "codeforces_url": ""},\n'
            '  "headline": "",\n'
            '  "summary": "",\n'
            '  "career_objective": "",\n'
            '  "education": [{"degree": "", "specialization": "", "institution": "", "start_date": "", "end_date": "", "cgpa": "", "description": ""}],\n'
            '  "experience": [{"title": "", "company": "", "location": "", "start_date": "", "end_date": "", "responsibilities": "", "achievements": ""}],\n'
            '  "projects": [{"project_name": "", "description": "", "technologies": "", "github_url": "", "live_url": "", "start_date": "", "end_date": ""}],\n'
            '  "skills": {"programming_languages": [], "frameworks": [], "databases": [], "cloud_technologies": [], "devops_tools": [], "ai_ml_technologies": [], "soft_skills": []},\n'
            '  "certifications": [{"certification_name": "", "issuer": "", "issue_date": "", "expiry_date": "", "credential_url": ""}],\n'
            '  "job_preferences": {"preferred_roles": "", "preferred_locations": "", "work_mode": "", "expected_salary_min": null, "expected_salary_max": null, "years_of_experience": ""}\n'
            "}"
        )

    def validate_file(self, filename: str, content: bytes, content_type: str = "", max_size_bytes: int = 10 * 1024 * 1024) -> str:
        if not filename:
            raise ResumeFileValidationError("No filename was provided.")
        if "\x00" in filename or ".." in filename or filename.replace("\\", "/").startswith("/"):
            raise ResumeFileValidationError("Unsafe filename rejected.")
        if not content:
            raise ResumeFileValidationError("Uploaded file is empty.")
        if len(content) > max_size_bytes:
            raise ResumeFileValidationError(f"File exceeds max size of {max_size_bytes // (1024 * 1024)} MB.")
        ext = Path(filename).suffix.lower()
        if ext not in self.allowed_extensions:
            raise ResumeFileValidationError(f"Unsupported file type: {ext or 'no extension'}. Upload PDF, DOCX, DOC, or TXT.")
        detected = self._detect_mime_type(content, ext)
        expected = self.allowed_mime_types.get(ext, set())
        submitted = (content_type or "").split(";")[0].strip().lower()
        if submitted and submitted not in expected and detected not in expected:
            raise ResumeFileValidationError(f"File content does not match extension {ext}. Upload a valid {ext.upper()} file.")
        return ext

    def ingest(self, content: bytes, filename: str, content_type: str = "", max_size_bytes: int = 10 * 1024 * 1024) -> Dict[str, Any]:
        ext = self.validate_file(filename, content, content_type, max_size_bytes)
        logger.info("resume_ingest_start filename=%s ext=%s size=%s", filename, ext, len(content))
        raw_text = ""
        cleaned_text = ""
        text_source = "none"
        try:
            raw_text = self.extract_text(content, filename)
            text_source = self.last_extraction_metadata.get("extraction_method", "unknown")
            cleaned_text = self.clean_text(raw_text)
            logger.info("resume_ingest_extracted filename=%s method=%s chars=%s", filename, text_source, len(cleaned_text))
        except Exception as exc:
            logger.error("resume_ingest_extraction_failed filename=%s error=%s", filename, exc, exc_info=True)
            raw_text = content.decode("utf-8", errors="ignore") if content else ""
            cleaned_text = self.clean_text(raw_text)
            text_source = f"raw_fallback_after_error:{type(exc).__name__}"
            logger.info("resume_ingest_raw_fallback filename=%s fallback_chars=%s", filename, len(cleaned_text))

        if not cleaned_text:
            logger.error("resume_ingest_no_text filename=%s ext=%s", filename, ext)
            raise ResumeIngestionError("Could not extract readable text from the uploaded resume. Ensure the file is not a scanned image-only PDF, and try a DOCX or TXT file instead.")

        try:
            parsed, parser_warning = self._parse_or_fallback(cleaned_text)
            logger.info("resume_ingest_parsed filename=%s warning=%s", filename, parser_warning)
        except Exception as exc:
            logger.error("resume_ingest_parse_failed filename=%s error=%s", filename, exc, exc_info=True)
            parsed = self._rule_based_extract(cleaned_text, f"Parsing failed: {exc}")
            parser_warning = "Both LLM and rule-based parsing failed; using minimal extraction."

        metadata = {
            "filename": Path(filename).name,
            "extension": ext,
            "text_chars": len(cleaned_text),
            "parser_model": self.model,
            "parser_warning": parser_warning,
            "extraction_method": text_source,
            **self.last_extraction_metadata,
        }
        return {
            "raw_text": cleaned_text,
            "structured_resume": parsed,
            "parsed": parsed,
            "metadata": metadata,
        }

    def extract_text(self, content: bytes, filename: str) -> str:
        suffix = Path(filename).suffix.lower()
        self.last_extraction_metadata = {"extraction_method": suffix.lstrip(".") or "unknown"}
        logger.info("extract_text_start filename=%s suffix=%s", filename, suffix)
        if suffix == ".pdf":
            return self._extract_pdf(content)
        if suffix == ".docx":
            return self._extract_docx(content)
        if suffix == ".txt":
            return self._extract_txt(content)
        if suffix == ".doc":
            raise ResumeIngestionError("Legacy .doc files are accepted but cannot be parsed reliably by this server. Convert the resume to DOCX, PDF, or TXT and upload again.")
        raise ResumeIngestionError(f"Unsupported file type: {suffix or 'no extension'}. Upload PDF, DOCX, DOC, or TXT.")

    def clean_text(self, text: str) -> str:
        text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n")]
        lines = [line for line in lines if line]
        cleaned = []
        previous_blank = False
        for line in lines:
            if not line:
                if not previous_blank:
                    cleaned.append("")
                previous_blank = True
            else:
                cleaned.append(line)
                previous_blank = False
        return "\n".join(cleaned).strip()[: self.max_text_chars]

    def _parse_or_fallback(self, cleaned_text: str) -> Tuple[Dict[str, Any], str]:
        try:
            return self.parse(cleaned_text), ""
        except ResumeIngestionError as exc:
            return self._rule_based_extract(cleaned_text, str(exc)), "LLM parsing unavailable; deterministic fallback was used for review."

    def parse(self, resume_text: str) -> Dict[str, Any]:
        prompt = (
            f"{self.system_prompt}\n\n"
            "Resume text:\n"
            "--------------------\n"
            f"{resume_text}\n"
            "--------------------\n"
            "Return ONLY the JSON object matching the schema."
        )
        try:
            raw = self.provider.generate(
                [{"role": "user", "content": prompt}],
                timeout=120,
                retries=1,
            )
            data = self._extract_json(raw)
            return self._normalize(data)
        except Exception as exc:
            raise ResumeIngestionError("Resume parser service temporarily unavailable. Please try again later.") from exc

    def _extract_pdf(self, content: bytes) -> str:
        errors = []
        extractors = [
            ("pdfplumber", self._extract_pdf_pdfplumber),
            ("pymupdf", self._extract_pdf_pymupdf),
            ("pypdf", self._extract_pdf_pypdf),
        ]
        for name, extractor in extractors:
            try:
                logger.info("pdf_extract_try extractor=%s", name)
                text = extractor(content)
                if text and text.strip():
                    logger.info("pdf_extract_success extractor=%s chars=%s", name, len(text))
                    self.last_extraction_metadata = {"extraction_method": name, "ocr_used": False}
                    return text
                logger.warning("pdf_extract_empty extractor=%s", name)
                errors.append(f"{name}: no text")
            except Exception as exc:
                if self._is_password_error(exc):
                    logger.error("pdf_extract_password_protected")
                    raise ResumeIngestionError("PDF is password protected. Remove the password and upload again.") from exc
                logger.warning("pdf_extract_error extractor=%s error=%s", name, exc)
                errors.append(f"{name}: {exc}")
        logger.info("pdf_extract_ocr_fallback errors=%s", errors)
        ocr_text = self._ocr_pdf(content)
        if ocr_text and ocr_text.strip():
            logger.info("pdf_extract_ocr_success chars=%s", len(ocr_text))
            self.last_extraction_metadata = {"extraction_method": "ocr", "ocr_used": True}
            return ocr_text
        logger.error("pdf_extract_all_failed errors=%s", errors)
        raise ResumeIngestionError("PDF extraction failed. Upload a text-based PDF or enable OCR support on the server.")

    def _extract_pdf_pdfplumber(self, content: bytes) -> str:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)

    def _extract_pdf_pymupdf(self, content: bytes) -> str:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as doc:
            return "\n".join(page.get_text("text") or "" for page in doc)

    def _extract_pdf_pypdf(self, content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages)

    def _ocr_pdf(self, content: bytes) -> str:
        try:
            import fitz
            import pytesseract
            from PIL import Image
        except Exception:
            return ""
        parts = []
        try:
            with fitz.open(stream=content, filetype="pdf") as doc:
                for page in doc[: min(len(doc), 10)]:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    image = Image.open(io.BytesIO(pix.tobytes("png")))
                    text = pytesseract.image_to_string(image)
                    if text and text.strip():
                        parts.append(text)
        except Exception:
            return ""
        return "\n".join(parts)

    def _extract_docx(self, content: bytes) -> str:
        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            for section in doc.sections:
                for paragraph in section.header.paragraphs + section.footer.paragraphs:
                    if paragraph.text:
                        parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells if cell.text)
                    if row_text:
                        parts.append(row_text)
            result = "\n".join(parts)
            logger.info("docx_extract_success chars=%s", len(result))
            return result
        except Exception as exc:
            logger.error("docx_extract_error error=%s", exc, exc_info=True)
            raise ResumeIngestionError("DOCX extraction failed. The file may be corrupted or not a valid DOCX.") from exc

    def _extract_txt(self, content: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1252"):
            try:
                text = content.decode(encoding)
                logger.info("txt_extract_success encoding=%s chars=%s", encoding, len(text))
                return text
            except UnicodeDecodeError:
                continue
        text = content.decode("utf-8", errors="ignore")
        logger.info("txt_extract_fallback chars=%s", len(text))
        return text

    def _detect_mime_type(self, content: bytes, ext: str) -> str:
        if ext == ".pdf" and content.startswith(b"%PDF"):
            return "application/pdf"
        if ext == ".docx" and content[:2] == b"PK":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == ".doc" and content[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
            return "application/msword"
        if ext == ".txt":
            try:
                content.decode("utf-8")
                return "text/plain"
            except UnicodeDecodeError:
                return "application/octet-stream"
        return ""

    def _is_password_error(self, exc: Exception) -> bool:
        message = str(exc).lower()
        return "encrypt" in message or "password" in message or "permission" in message

    def _extract_json(self, raw: str) -> Dict[str, Any]:
        text = raw.strip()
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\n?", "", text)
            text = re.sub(r"\n?```$", "", text)
        start = text.find("{")
        end = text.rfind("}")
        if start == -1 or end == -1 or end < start:
            raise ValueError("LLM response did not contain a JSON object")
        json_str = text[start : end + 1]
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as exc:
            raise ValueError("LLM response was not valid JSON") from exc

    def _normalize(self, data: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(data, dict):
            raise ValueError("Parsed output is not a JSON object")
        info = data.get("personal_info") or {}
        if not isinstance(info, dict):
            info = {}
        skills = data.get("skills") or {}
        if isinstance(skills, list):
            skills = {"general": skills}
        if not isinstance(skills, dict):
            skills = {}

        personal_confidence = {"full_name": 0.95, "email": 0.98, "phone": 0.92, "location": 0.85, "linkedin_url": 0.90, "github_url": 0.90, "portfolio_url": 0.85, "leetcode_url": 0.88, "hackerrank_url": 0.88, "codeforces_url": 0.88}
        
        return {
            "personal_info": {
                "full_name": ConfidenceScorer.scored(_clean_string(info.get("full_name") or info.get("name")), personal_confidence["full_name"]),
                "email": ConfidenceScorer.scored(_clean_string(info.get("email")), personal_confidence["email"]),
                "phone": ConfidenceScorer.scored(_clean_string(info.get("phone")), personal_confidence["phone"]),
                "location": ConfidenceScorer.scored(_clean_string(info.get("location")), personal_confidence["location"]),
                "linkedin_url": ConfidenceScorer.scored(_clean_url(info.get("linkedin_url") or info.get("linkedin")), personal_confidence["linkedin_url"]),
                "github_url": ConfidenceScorer.scored(_clean_url(info.get("github_url") or info.get("github")), personal_confidence["github_url"]),
                "portfolio_url": ConfidenceScorer.scored(_clean_url(info.get("portfolio_url") or info.get("portfolio")), personal_confidence["portfolio_url"]),
                "leetcode_url": ConfidenceScorer.scored(_clean_url(info.get("leetcode_url")), personal_confidence["leetcode_url"]),
                "hackerrank_url": ConfidenceScorer.scored(_clean_url(info.get("hackerrank_url")), personal_confidence["hackerrank_url"]),
                "codeforces_url": ConfidenceScorer.scored(_clean_url(info.get("codeforces_url")), personal_confidence["codeforces_url"]),
            },
            "headline": ConfidenceScorer.scored(_clean_string(data.get("headline")), 0.90),
            "summary": ConfidenceScorer.scored(_clean_string(data.get("summary")), 0.88),
            "career_objective": ConfidenceScorer.scored(_clean_string(data.get("career_objective")), 0.85),
            "education": [_scored_education(item) for item in _as_list(data.get("education"))],
            "experience": [_scored_experience(item) for item in _as_list(data.get("experience"))],
            "projects": [_scored_project(item) for item in _as_list(data.get("projects"))],
            "skills": {str(key): _as_string_list(value) for key, value in skills.items()},
            "certifications": [_scored_certification(item) for item in _as_list(data.get("certifications"))],
            "job_preferences": _normalize_dict(data.get("job_preferences") or {}),
        }

    # ------------------------------------------------------------------
    # Rule-based fallback extractor (used when LLM is unavailable/fails)
    # ------------------------------------------------------------------
    def _rule_based_extract(self, text: str, warning: str) -> Dict[str, Any]:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        full_text = "\n".join(lines)

        classifier = SectionClassifier()
        sections = classifier.classify(lines)

        def section_lines(section_name: str) -> List[str]:
            indices = sections.get(section_name, [])
            return [lines[i] for i in indices]

        personal = self._extract_personal_info(section_lines("header") or lines[:5], full_text)
        education = self._extract_education(section_lines("education"))
        experience = self._extract_experience(section_lines("experience"))
        projects = self._extract_projects(section_lines("projects"))
        skills_raw = self._extract_skills(section_lines("skills") or lines)
        categorized_skills = self._categorize_skills(skills_raw)
        certifications = self._extract_certifications(section_lines("certifications"))

        headline = self._extract_headline(lines, personal)
        summary = self._build_summary(lines)

        return self._normalize({
            "personal_info": personal,
            "headline": headline,
            "summary": summary,
            "career_objective": "",
            "education": education,
            "experience": experience,
            "projects": projects,
            "skills": categorized_skills,
            "certifications": certifications,
            "job_preferences": {},
        }) | {"parser_warning": f"Rule-based extraction used: {warning}"}

    def _extract_personal_info(self, lines: List[str], full_text: str) -> Dict[str, str]:
        info: Dict[str, str] = {
            "full_name": "", "email": "", "phone": "", "location": "",
            "linkedin_url": "", "github_url": "", "portfolio_url": "", "leetcode_url": "", "hackerrank_url": "", "codeforces_url": ""
        }
        email_re = r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
        phone_re = r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}"
        linkedin_re = r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+/?"
        github_re = r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_.-]+/?"
        leetcode_re = r"(?:https?://)?(?:www\.)?leetcode\.com/[A-Za-z0-9_-]+/?"
        hackerrank_re = r"(?:https?://)?(?:www\.)?hackerrank\.com/[A-Za-z0-9_-]+/?"
        codeforces_re = r"(?:https?://)?(?:www\.)?codeforces\.com/[A-Za-z0-9_-]+/?"
        portfolio_re = r"(?:https?://)?(?:www\.)?[A-Za-z0-9_-]+\.(?:vercel\.app|netlify\.app|herokuapp\.com|github\.io|me|dev)/[A-Za-z0-9_/-]*"

        for line in lines[:5]:
            if not info["email"]:
                m = re.search(email_re, line)
                if m:
                    info["email"] = m.group(0)
            if not info["phone"]:
                m = re.search(phone_re, line)
                if m:
                    info["phone"] = m.group(0)
            if not info["linkedin_url"]:
                m = re.search(linkedin_re, line)
                if m:
                    info["linkedin_url"] = _clean_url(m.group(0))
            if not info["github_url"]:
                m = re.search(github_re, line)
                if m:
                    info["github_url"] = _clean_url(m.group(0))
            if not info["portfolio_url"]:
                m = re.search(portfolio_re, line, re.IGNORECASE)
                if m:
                    info["portfolio_url"] = _clean_url(m.group(0))
            if not info["leetcode_url"]:
                m = re.search(leetcode_re, line, re.IGNORECASE)
                if m:
                    info["leetcode_url"] = _clean_url(m.group(0))
            if not info["hackerrank_url"]:
                m = re.search(hackerrank_re, line, re.IGNORECASE)
                if m:
                    info["hackerrank_url"] = _clean_url(m.group(0))
            if not info["codeforces_url"]:
                m = re.search(codeforces_re, line, re.IGNORECASE)
                if m:
                    info["codeforces_url"] = _clean_url(m.group(0))

        for line in full_text.splitlines():
            if not info["leetcode_url"]:
                m = re.search(leetcode_re, line, re.IGNORECASE)
                if m:
                    info["leetcode_url"] = _clean_url(m.group(0))
            if not info["hackerrank_url"]:
                m = re.search(hackerrank_re, line, re.IGNORECASE)
                if m:
                    info["hackerrank_url"] = _clean_url(m.group(0))
            if not info["codeforces_url"]:
                m = re.search(codeforces_re, line, re.IGNORECASE)
                if m:
                    info["codeforces_url"] = _clean_url(m.group(0))

        # Name: first non-empty line that is not contact info
        for line in lines:
            if not line:
                continue
            if re.search(email_re, line) or re.search(phone_re, line):
                continue
            if "linkedin" in line.lower() or "github" in line.lower():
                continue
            if len(line) < 60 and not any(kw in line.lower() for kw in ["university", "college", "skills", "experience", "education", "project", "summary", "objective", "profile", "career", "internship", "employment"]):
                info["full_name"] = line
                break

        # Location extraction with proper city/state/country format
        state_abbreviations = {"AL","AK","AZ","AR","CA","CO","CT","DE","FL","GA","HI","ID","IL","IN","IA","KS","KY","LA","ME","MD","MA","MI","MN","MS","MO","MT","NE","NV","NH","NJ","NM","NY","NC","ND","OH","OK","OR","PA","RI","SC","SD","TN","TX","UT","VT","VA","WA","WV","WI","WY"}
        country_names = {"USA","US","United States","United States of America","India","Indonesia","UK","United Kingdom","Canada","Australia","Germany","Netherlands","Singapore","Japan","China","Remote","Hybrid"}

        for line in lines[:5]:
            if info["location"]:
                break
            location_matches = re.findall(r"([A-Za-z][A-Za-z\s]+(?:,\s*[A-Za-z\s]+){1,2})", line)
            for match in location_matches:
                parts = [p.strip() for p in match.split(",")]
                valid_parts = []
                for part in parts:
                    if part.strip().lower() in INDIAN_CITIES or part.strip().lower() in US_CITIES:
                        valid_parts.append(part.strip())
                    elif part.strip() in state_abbreviations or part.strip().lower() in country_names:
                        valid_parts.append(part.strip())
                    elif len(part.strip()) <= 20 and re.match(r'^[A-Za-z\s]+$', part.strip()):
                        for city in INDIAN_CITIES | US_CITIES:
                            if city in part.strip().lower():
                                valid_parts.append(part.strip())
                                break
                if len(valid_parts) >= 2:
                    info["location"] = ", ".join(valid_parts[:3])
                    break
                elif len(valid_parts) == 1:
                    info["location"] = valid_parts[0]

        if not info["location"]:
            for line in lines[:5]:
                for city in INDIAN_CITIES | US_CITIES:
                    pattern = r'\b' + re.escape(city) + r'\b'
                    if re.search(pattern, line, re.IGNORECASE):
                        info["location"] = city.title()
                        if "India" in line or "USA" in line or "US" in line or "United States" in line:
                            if "India" in line:
                                info["location"] = f"{info['location']}, India"
                            elif "USA" in line or "US" in line:
                                info["location"] = f"{info['location']}, USA"
                        break
                if info["location"]:
                    break

        return info

    def _extract_headline(self, lines: List[str], personal: Dict[str, str]) -> str:
        name = personal.get("full_name", "")
        for i, line in enumerate(lines):
            if line == name:
                if i + 1 < len(lines):
                    candidate = lines[i + 1]
                    if candidate and not re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", candidate) and not re.search(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}", candidate):
                        return candidate
                break
        return ""

    def _build_summary(self, lines: List[str]) -> str:
        skip_patterns = [
            r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}",
            r"(?:\+?\d{1,3}[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}",
            r"linkedin\.com/in/[A-Za-z0-9_-]+/?",
            r"github\.com/[A-Za-z0-9_.-]+/?",
            r"^(skills?|education|experience|projects?|certifications?|summary|objective|profile|contact|references?)\s*:?\s*$",
            r"^skills?\s*:",
        ]
        summary_lines = []
        for line in lines:
            if any(re.search(p, line, re.IGNORECASE) for p in skip_patterns):
                continue
            if re.match(r"^[A-Za-z\s]+,\s*[A-Za-z]{2}$", line.strip()):
                continue
            if len(line) > 5:
                summary_lines.append(line)
        summary = " ".join(summary_lines)
        sentences = re.split(r'(?<=[.!?])\s+', summary)
        return " ".join(sentences[:3])[:500]

    def _extract_education(self, lines: List[str]) -> List[Dict[str, str]]:
        results = []
        degree_patterns = [
            (r"\b(B\.?Tech|B\.E\.?|Bachelor\s+of\s+Technology|Bachelor\s+of\s+Engineering)\b(?:[\s,]+([^\d,]+))?", "Bachelor of Technology"),
            (r"\b(BSc|B\.Sc|BS|B\.S|Bachelor\s+of\s+Science)\b(?:[\s,]*\s+in\s+)?([^\d,]+)?", "Bachelor of Science"),
            (r"\b(Bachelor\s+of\s+Arts|B\.A|BA)\b(?:[\s,]+([^\d,]+))?", "Bachelor of Arts"),
            (r"\b(Bachelor\s+of\s+Commerce|B\.Com)\b(?:[\s,]+([^\d,]+))?", "Bachelor of Commerce"),
            (r"\b(BCA)\b(?:[\s,]+([^\d,]+))?", "Bachelor of Computer Applications"),
            (r"\b(M\.?Tech|M\.E\.?|Master\s+of\s+Technology|Master\s+of\s+Engineering)\b(?:[\s,]+([^\d,]+))?", "Master of Technology"),
            (r"\b(MSc|M\.Sc|MS|M\.S|Master\s+of\s+Science)\b(?:[\s,]*\s+in\s+)?([^\d,]+)?", "Master of Science"),
            (r"\b(MBA|Master\s+of\s+Business\s+Administration)\b(?:[\s,]+([^\d,]+))?", "Master of Business Administration"),
            (r"\b(MCA|Master\s+of\s+Computer\s+Applications)\b(?:[\s,]+([^\d,]+))?", "Master of Computer Applications"),
            (r"\b(Master\s+of\s+Arts|M\.A|MA)\b(?:[\s,]+([^\d,]+))?", "Master of Arts"),
            (r"\b(PhD|Ph\.D|Doctor\s+of\s+Philosophy)\b(?:[\s,]+([^\d,]+))?", "Doctor of Philosophy"),
            (r"\b(Diploma)\b(?:[\s,]+([^\d,]+))?", "Diploma"),
        ]
        year_re = r"(?:19|20)\d{2}"
        cgpa_re = r"(?:CGPA|GPA|CPI)[\s:]*([0-9](?:\.[0-9])?)"
        institution_keywords = r"(?:University|College|Institute|School|Academy|IIT|NIT|VIT|MIT|Bits|IIIT|College of Engineering)"

        current_entry: Dict[str, str] = {}
        in_education = False

        for i, line in enumerate(lines):
            lower = line.lower()
            if re.search(r'\b(education|academic|academics|qualification|qualifications)\b', lower):
                in_education = True
                continue
            if SectionClassifier.is_boundary(line):
                in_education = False
                if current_entry:
                    results.append(current_entry)
                    current_entry = {}
                continue

            if not in_education:
                continue

            found_degree = False
            for pattern, default_degree in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    if current_entry:
                        results.append(current_entry)
                    current_entry = {"degree": "", "specialization": "", "institution": "", "start_date": "", "end_date": "", "cgpa": "", "description": ""}
                    current_entry["degree"] = default_degree

                    if match.lastindex and match.group(match.lastindex):
                        specialization = match.group(match.lastindex).strip()
                        if specialization and len(specialization) > 2:
                            current_entry["specialization"] = _clean_string(specialization)

                    after_degree = line[match.end():]
                    inst_m = re.search(institution_keywords, after_degree, re.IGNORECASE)
                    if inst_m:
                        current_entry["institution"] = _clean_string(after_degree[inst_m.start():])
                    found_degree = True
                    break

            if found_degree:
                continue

            if current_entry and re.search(institution_keywords, line, re.IGNORECASE):
                if not current_entry["institution"]:
                    current_entry["institution"] = _clean_string(line)

            cgpa_m = re.search(cgpa_re, line, re.IGNORECASE)
            if cgpa_m:
                current_entry["cgpa"] = cgpa_m.group(1)

            years = re.findall(year_re, line)
            if len(years) >= 2:
                current_entry["start_date"] = years[0]
                current_entry["end_date"] = years[-1]
            elif len(years) == 1 and not current_entry.get("end_date"):
                current_entry["end_date"] = years[0]

        if current_entry:
            results.append(current_entry)

        for entry in results:
            if entry.get("institution"):
                inst = entry["institution"]
                inst = re.sub(r',\s*(?:19|20)\d{2}.*$', '', inst)
                inst = re.sub(r',\s*CGPA[\s:]*[\d.]+.*$', '', inst, flags=re.IGNORECASE)
                inst = re.sub(r',\s*GPA[\s:]*[\d.]+.*$', '', inst, flags=re.IGNORECASE)
                inst = re.sub(r',\s*(?:Jan(?:uary)?|Feb(?:uary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s*(?:19|20)\d{2}.*$', '', inst, flags=re.IGNORECASE)
                entry["institution"] = inst.strip(" ,-")

        return results if results else []

    def _extract_experience(self, lines: List[str]) -> List[Dict[str, str]]:
        results = []
        title_keywords = r"(?:Engineer|Developer|SDE|Manager|Analyst|Designer|Architect|Consultant|Intern|Lead|Director|VP|Head|Specialist|Coordinator|Associate|Executive|Officer|Scientist|Researcher|Advisor|Strategist)"
        company_keywords = r"(?:Inc|LLC|Ltd|Corp|Company|Corporation|Technologies|Solutions|Systems|Services|Labs|Startup|Amazon|Google|Meta|Microsoft|Apple|Netflix|Uber|Airbnb|Spotify|Adobe|Oracle|Salesforce|IBM|Intel|AMD|NVIDIA|Tesla|SpaceX|Flipkart|Zomato|Swiggy|Paytm|Ola|Reliance|TCS|Infosys|Wipro|HCL|TechMahindra|Accenture|Deloitte|KPMG|EY|PwC)"

        current_entry: Dict[str, str] = {}
        in_exp = False

        for i, line in enumerate(lines):
            lower = line.lower()
            if re.search(r'\b(experience|work|employment|professional|internship)\b', lower) and not current_entry:
                in_exp = True
                continue
            if re.search(r'\b(education|project|skill|certification|summary|objective)\b', lower) and in_exp:
                in_exp = False
                if current_entry:
                    results.append(current_entry)
                    current_entry = {}
                continue

            if current_entry and not in_exp:
                results.append(current_entry)
                current_entry = {}

            if not in_exp:
                continue

            role_company_m = re.search(r"^(.+?)\s+at\s+(.+?)(?:,\s*(.+))?$", line, re.IGNORECASE)
            if role_company_m:
                if current_entry:
                    results.append(current_entry)
                current_entry = {"title": "", "company": "", "location": "", "start_date": "", "end_date": "", "responsibilities": "", "achievements": ""}
                current_entry["title"] = _clean_string(role_company_m.group(1))
                current_entry["company"] = _clean_string(role_company_m.group(2))
                if role_company_m.group(3):
                    dates = role_company_m.group(3).strip()
                    years = re.findall(r"(?:19|20)\d{2}", dates)
                    if years:
                        current_entry["start_date"] = years[0]
                        current_entry["end_date"] = years[-1] if len(years) > 1 else ""
                continue

            if re.search(company_keywords, line, re.IGNORECASE) and not current_entry.get("company"):
                current_entry["company"] = _clean_string(line)
                continue

            if re.search(r'^' + title_keywords, line, re.IGNORECASE) and not current_entry.get("title"):
                current_entry["title"] = _clean_string(line)
                continue

            dates_m = re.search(r"(?:19|20)\d{2}\s*[-–—]\s*(?:Present|(?:19|20)\d{2})", line, re.IGNORECASE)
            if dates_m:
                years = re.findall(r"(?:19|20)\d{2}", dates_m.group(0))
                if len(years) >= 2:
                    current_entry["start_date"] = years[0]
                    current_entry["end_date"] = years[1]
                continue

            if current_entry and line.startswith(("-", "*", "\u2022", "\u00b7")):
                bullet = line.lstrip("-*\u2022\u00b7 ").strip()
                if bullet:
                    if not current_entry.get("responsibilities"):
                        current_entry["responsibilities"] = bullet
                    else:
                        current_entry["responsibilities"] += "\n" + bullet
            elif current_entry and not re.search(r"^\s*(?:19|20)\d{2}\s*[-–—]\s*(?:Present|(?:19|20)\d{2})\s*$", line, re.IGNORECASE) and not re.search(r"^(?:\w[\w\s]+)\s+at\s+(?:\w[\w\s]+)", line, re.IGNORECASE) and not re.search(r"^(?:Inc|LLC|Ltd|Corp|Company|Corporation|Technologies|Solutions|Systems|Services|Labs|Startup)", line, re.IGNORECASE) and not re.search(r"^(?:B\.?Tech|BE|BSc|M\.?Tech|MTech|MBA|PhD|BS|MS|Bachelor|Master)", line, re.IGNORECASE) and len(line) > 5:
                if not current_entry.get("responsibilities"):
                    current_entry["responsibilities"] = line
                else:
                    current_entry["responsibilities"] += "\n" + line

        if current_entry:
            results.append(current_entry)

        return results if results else []

    def _extract_projects(self, lines: List[str]) -> List[Dict[str, str]]:
        results = []
        current_entry: Dict[str, str] = {}

        for i, line in enumerate(lines):
            lower = line.lower()

            if SectionClassifier.is_boundary(line):
                if current_entry:
                    results.append(current_entry)
                    current_entry = {}
                continue

            if re.search(r'\b(certification|certificate|certified|credential|license|certifications?|certificates?)\b', lower):
                if current_entry:
                    results.append(current_entry)
                    current_entry = {}
                continue

            if re.search(r'^\s*(?:https?://)?(?:www\.)?(?:linkedin\.com|github\.com|leetcode\.com|hackerrank\.com|codeforces\.com)\.?$', lower):
                continue

            if re.search(r'^\s*(issued\s+by|from|issuer|credential\s+id|certificate\s+id)\s*:?\s*', lower):
                if current_entry:
                    results.append(current_entry)
                    current_entry = {}
                break

            if not current_entry:
                if line and not re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", line):
                    entry = {"project_name": _clean_string(line), "description": "", "technologies": "", "github_url": "", "live_url": "", "start_date": "", "end_date": ""}
                    parsed_first = self._parse_project_line(line)
                    entry.update(parsed_first)
                    current_entry = entry
                continue

            github_m = re.search(r"github\.com/[A-Za-z0-9_.-]+/?", line)
            if github_m:
                current_entry["github_url"] = _clean_url(github_m.group(0))
                continue

            live_m = re.search(r"https?://[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?", line)
            if live_m and "github" not in live_m.group(0).lower():
                current_entry["live_url"] = live_m.group(0)
                continue

            bullet_match = re.match(r"^[\d]+\.?\s*|-|\*|•", line)
            if bullet_match and current_entry.get("description"):
                current_entry["description"] += " " + line.strip()
            elif current_entry and line:
                if not current_entry.get("description"):
                    current_entry["description"] = line
                elif not current_entry.get("technologies"):
                    current_entry["technologies"] = line

        if current_entry:
            results.append(current_entry)

        return results if results else []

    def _parse_project_line(self, line: str) -> Dict[str, str]:
        result = {"project_name": "", "technologies": "", "github_url": "", "live_url": ""}
        parts = re.split(r'\s+[-–|]\s+', line)
        if len(parts) >= 1:
            result["project_name"] = _clean_string(parts[0])
        if len(parts) >= 2:
            tech_or_url = parts[1]
            if re.search(r"github\.com/[A-Za-z0-9_.-]+/?", tech_or_url, re.IGNORECASE):
                result["github_url"] = _clean_url(tech_or_url)
            elif re.search(r"https?://", tech_or_url, re.IGNORECASE):
                result["live_url"] = tech_or_url
            else:
                result["technologies"] = _clean_string(tech_or_url)
        if len(parts) >= 3:
            for part in parts[1:-1]:
                if re.search(r"github\.com/[A-Za-z0-9_.-]+/?", part, re.IGNORECASE):
                    result["github_url"] = _clean_url(part)
                elif re.search(r"https?://", part, re.IGNORECASE):
                    result["live_url"] = part
                elif not result["technologies"]:
                    result["technologies"] = _clean_string(part)
        if not result["project_name"]:
            result["project_name"] = _clean_string(line)
        return result

    def _extract_skills(self, lines: List[str]) -> List[str]:
        skills: List[str] = []
        skill_section = True
        for line in lines:
            stripped = line.strip()
            lower = stripped.lower()

            if SectionClassifier.is_boundary(stripped):
                skill_keywords = re.compile(r'^(skills?|technical[\s_]skills?|core[\s_]competenc(?:y|ies)|technologies|programming[\s_]languages|tools?\s*&?\s*technologies|proficiencies?)\s*:?\s*$', re.IGNORECASE)
                if skill_keywords.match(lower):
                    skill_section = True
                    if ':' in stripped:
                        inline_skills = re.split(r'[,;|/·•]+', stripped.split(':', 1)[-1].strip())
                        for item in inline_skills:
                            item = item.strip()
                            if item and len(item) > 1 and not re.search(r'[^\x00-\x7F]', item):
                                skills.append(item)
                else:
                    skill_section = False
                continue

            if not skill_section:
                continue

            if stripped:
                items = re.split(r'[,;|/·•]+', stripped)
                for item in items:
                    item = item.strip()
                    if item and len(item) > 1 and not re.search(r'[^\x00-\x7F]', item):
                        skills.append(item)
        return skills

    def _categorize_skills(self, skills: List[str]) -> Dict[str, List[str]]:
        categorized: Dict[str, List[str]] = {}
        uncategorized: List[str] = []
        for skill in skills:
            normalized = self._normalize_skill_name(skill)
            category = CANONICAL_SKILLS.get(normalized)
            if category:
                if category not in categorized:
                    categorized[category] = []
                if skill not in categorized[category]:
                    categorized[category].append(skill)
            else:
                if skill not in uncategorized:
                    uncategorized.append(skill)
        if uncategorized:
            categorized.setdefault("general", [])
            categorized["general"].extend(uncategorized)
        return categorized

    @staticmethod
    def _normalize_skill_name(skill: str) -> str:
        normalized = skill.lower().strip()
        aliases = {
            "js": "javascript", "ts": "typescript", "k8s": "kubernetes", "ml": "machine learning",
            "ai": "artificial intelligence", "devops": "ci/cd", "cicd": "ci/cd",
            "node": "node.js", "nodejs": "node.js", "node.js": "node.js", "reactjs": "react", "vuejs": "vue", "next": "next.js",
            "nextjs": "next.js", "expressjs": "express.js", "express.js": "express.js", "springboot": "spring boot",
            "spring-boot": "spring boot", "postgres": "postgresql", "mongo": "mongodb", "redis cache": "redis",
            "aws cloud": "aws", "azure cloud": "azure", "gcp cloud": "gcp",
            "docker container": "docker", "k8s orchestration": "kubernetes",
        }
        if normalized in aliases:
            return aliases[normalized]
        if "node" in normalized and "node.js" not in normalized:
            return "node.js"
        return normalized

    def _extract_certifications(self, lines: List[str]) -> List[Dict[str, str]]:
        results = []
        current: Dict[str, str] = {}
        cert_keywords = r"(?:certification|certificate|certified|credential|license)"
        cert_section_start = re.compile(r'^\s*(certifications?|certificates?|certified|credentials?|licenses?)\s*:?\s*$', re.IGNORECASE)
        section_end = re.compile(r'^\s*(education|experience|skill|project|summary|objective|achievement|award|honor|interest|hobby|language|declaration|projects?)\s*:?\s*$', re.IGNORECASE)

        in_cert = False
        for line in lines:
            lower = line.lower().strip()
            if cert_section_start.match(lower):
                in_cert = True
                continue
            if in_cert and section_end.match(lower):
                in_cert = False
                if current:
                    results.append(current)
                    current = {}
                continue
            if not in_cert:
                continue
            if not current:
                current = {"certification_name": _clean_string(line), "issuer": "", "issue_date": "", "expiry_date": "", "credential_url": ""}
                continue
            issuer_m = re.search(r"(?:by|from|issued by|offered by|from)\s+(.+)", line, re.IGNORECASE)
            if issuer_m:
                current["issuer"] = _clean_string(issuer_m.group(1))
                continue
            date_m = re.search(r"(?:19|20)\d{2}", line)
            if date_m:
                if not current["issue_date"]:
                    current["issue_date"] = date_m.group(0)
                elif not current["expiry_date"]:
                    current["expiry_date"] = date_m.group(0)

        if current:
            results.append(current)
        return results


class ResumeParser:
    def __init__(self, model: str = "gemini-2.0-flash"):
        self.pipeline = ResumeIngestionPipeline(model=model)

    def parse(self, resume_text: str) -> Dict[str, Any]:
        return self.pipeline.parse(resume_text)


def _clean_string(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()[:2000]


def _clean_url(value: Any) -> str:
    text = _clean_string(value)
    if not text:
        return ""
    if not re.match(r"^https?://", text, flags=re.IGNORECASE):
        return f"https://{text}"
    return text


def _first_match(pattern: str, text: str, flags: int = 0) -> str:
    match = re.search(pattern, text, flags=flags)
    return match.group(1) if match and match.lastindex else (match.group(0) if match else "")


def _as_list(value: Any) -> List[Any]:
    if value is None:
        return []
    if isinstance(value, list):
        return value
    return [value]


def _as_string_list(value: Any) -> List[str]:
    if value is None:
        return []
    if isinstance(value, str):
        return [item.strip() for item in re.split(r"[,;\n]", value) if item.strip()]
    if isinstance(value, list):
        return [_clean_string(item) for item in value if _clean_string(item)]
    return [_clean_string(value)] if _clean_string(value) else []


def _normalize_dict(value: Any) -> Dict[str, Any]:
    if isinstance(value, dict):
        return {str(key): _clean_string(val) for key, val in value.items()}
    return {}


def _scored_education(value: Any) -> Dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {
        "degree": ConfidenceScorer.scored(value.get("degree", ""), 0.90),
        "specialization": ConfidenceScorer.scored(value.get("specialization", ""), 0.85),
        "institution": ConfidenceScorer.scored(value.get("institution", ""), 0.92),
        "start_date": ConfidenceScorer.scored(value.get("start_date", ""), 0.80),
        "end_date": ConfidenceScorer.scored(value.get("end_date", ""), 0.80),
        "cgpa": ConfidenceScorer.scored(value.get("cgpa", ""), 0.88),
        "description": ConfidenceScorer.scored(value.get("description", ""), 0.75),
    }


def _scored_experience(value: Any) -> Dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {
        "title": ConfidenceScorer.scored(value.get("title", ""), 0.92),
        "company": ConfidenceScorer.scored(value.get("company", ""), 0.90),
        "location": ConfidenceScorer.scored(value.get("location", ""), 0.85),
        "start_date": ConfidenceScorer.scored(value.get("start_date", ""), 0.85),
        "end_date": ConfidenceScorer.scored(value.get("end_date", ""), 0.85),
        "responsibilities": ConfidenceScorer.scored(value.get("responsibilities", ""), 0.88),
        "achievements": ConfidenceScorer.scored(value.get("achievements", ""), 0.85),
    }


def _scored_project(value: Any) -> Dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {
        "project_name": ConfidenceScorer.scored(value.get("project_name", ""), 0.92),
        "description": ConfidenceScorer.scored(value.get("description", ""), 0.88),
        "technologies": ConfidenceScorer.scored(value.get("technologies", ""), 0.85),
        "github_url": ConfidenceScorer.scored(value.get("github_url", ""), 0.80),
        "live_url": ConfidenceScorer.scored(value.get("live_url", ""), 0.80),
        "start_date": ConfidenceScorer.scored(value.get("start_date", ""), 0.75),
        "end_date": ConfidenceScorer.scored(value.get("end_date", ""), 0.75),
    }


def _scored_certification(value: Any) -> Dict[str, Any]:
    if not isinstance(value, dict):
        return {}
    return {
        "certification_name": ConfidenceScorer.scored(value.get("certification_name", ""), 0.90),
        "issuer": ConfidenceScorer.scored(value.get("issuer", ""), 0.80),
        "issue_date": ConfidenceScorer.scored(value.get("issue_date", ""), 0.85),
        "expiry_date": ConfidenceScorer.scored(value.get("expiry_date", ""), 0.85),
        "credential_url": ConfidenceScorer.scored(value.get("credential_url", ""), 0.75),
    }
